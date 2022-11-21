#Automaticaly recalculate=true
#Single model=false
import docx
import pandas as pd
import numpy as np
import os
import re
import win32com.client as win32
from win32com.client import constants

## скрипт конвертирует информацию из doc файлов в excel для каждой скважины отдельно
def save_as_docx(path):

    """Конвертация doc в docx"""
    # Opening MS Word
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(path)
    doc.Activate ()

    # Rename path with .docx
    new_file_abs = os.path.abspath(path)
    new_file_abs = re.sub(r'\.\w+$', '.DOCX', new_file_abs)

    # Save and Close
    word.ActiveDocument.SaveAs(
        new_file_abs, FileFormat=constants.wdFormatXMLDocument
    )
    doc.Close(False)

##Getting the original data from the document to a list
##
def Table_from_doc(filename, tab_number):
    """Функция для парсинга таблицы в формат Dataframe"""
    doc = docx.Document(filename)
    df = pd.DataFrame()
    tables = doc.tables[tab_number]
    column=int(tables._column_count)
    ls =[]
    for row in tables.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # print(paragraph.text)
                text = paragraph.text.replace("*","")
                ls.append(text)
    row=int(len(ls)/column-1)
    colname = ls[0:column]
    df = pd.DataFrame(np.array(ls[column:]).reshape(column, row))  #reshape to the table shape
    df = df.T
    df.columns=colname
    df.dropna(how='all')
    return df.iloc[:-1,]
##
def parse_extra_param(filename, name):
    """Собирает данные по альтитуде и магнитным поправкам в вордовских файлах
    с инклинометрией
    :arg1: TODO
    :returns: TODO

    """
    doc = docx.Document(filename)
    extra_info = [name]
    for n, x in enumerate(doc.paragraphs):
        if 'Альтитуда' in x.text:
            splitted = re.sub('([А-Я]+)', r' \1', re.sub('([А-Я]+)', r' \1', x.text)).split()
            extra_info.append(splitted[2])
            extra_info.append(splitted[5])
            extra_info.append(splitted[8])
    return extra_info
#

##
folder = os.getcwd()
extra = []

for root, dirs, files in os.walk(folder):
    for file in files:
        if file.endswith('DOC') and not file.startswith('~') and not file.endswith('DOCX'):
            save_as_docx(os.path.join(root, file))
    for file in files:
        if file.endswith('DOCX') and not file.startswith('~'):
            name = file.split("_")[0]+".xlsx"
            Table_from_doc(os.path.join(root, file),1).to_excel(name)
            extra.append(parse_extra_param(os.path.join(root, file),file.split("_")[0]))

print(extra)
df = pd.DataFrame(extra)  #reshape to the table shape
df.columns = ['WELL', 'Магнитная поправка', 'Альтитуда', 'Забой']
df.to_excel('extrainfo.xlsx')
            # paths.append(os.path.join(root, file))
##

#frame = Table_from_doc("32702_1_906.DOCX", 1)
#frame.to_excel('test.xlsx')

#doc = docx.Document("32702_1_906.DOCX")
#extra_info = []
#name = doc.paragraphs[18].text
#splitted = re.sub('([А-Я]+)', r' \1', re.sub('([А-Я]+)', r' \1', name)).split()
###
#for n, x in enumerate(doc.paragraphs):
#    if 'Альтитуда' in x.text:
#        splitted = re.sub('([А-Я]+)', r' \1', re.sub('([А-Я]+)', r' \1', x.text)).split()
#        extra_info.append(splitted[2])
#        extra_info.append(splitted[5])
#        extra_info.append(splitted[8])
#        print(extra_info)
###

#dir(doc.paragraphs[0])
